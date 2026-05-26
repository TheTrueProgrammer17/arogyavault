from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Add light gray background
    shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shading_elm)
    return p

doc = Document()

# Title Page
title = doc.add_heading('ArogyaVault Technical Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')
subtitle = doc.add_paragraph('Comprehensive Analysis and System Architecture (With Code Snippets)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 1. Project Overview
add_heading(doc, '1. Project Overview', 1)
doc.add_paragraph("ArogyaVault (also known as Medivault) is a modern, government-grade healthcare document management platform. It allows citizens to securely store their medical records and provides an emergency access system for first responders. Additionally, the platform features a highly secure Doctor Portal with a strict verification workflow. Doctors are granted temporary, key-based access to specific medical documents, ensuring robust data privacy and security.")

# 2. Features Implemented
add_heading(doc, '2. Features Implemented', 1)
doc.add_paragraph("• Secure User and Doctor Authentication with Role-Based Routing")
doc.add_paragraph("• Patient Dashboard for Medical Document Management")
doc.add_paragraph("• Family Dashboard for linking and managing dependents' accounts")
doc.add_paragraph("• Emergency QR Workflow for First Responders")
doc.add_paragraph("• Temporary Verification Key Generation (AV-XXXX-XXXX) with Expiry Timers")
doc.add_paragraph("• Doctor Verification Hub")

# 3. Folder Structure Explanation
add_heading(doc, '3. Folder Structure Explanation', 1)
doc.add_paragraph("The workspace follows a standard Vite + React modular structure:")
doc.add_paragraph("• /src/components: Reusable UI elements")
doc.add_paragraph("• /src/contexts: State management (AuthContext.jsx)")
doc.add_paragraph("• /src/pages: Application routes")
doc.add_paragraph("• /src/firebase.js: Firebase config")
doc.add_paragraph("• /src/App.jsx: Core application routing")

# 4. Frontend Technologies Used
add_heading(doc, '4. Frontend Technologies Used', 1)
doc.add_paragraph("• React 19")
doc.add_paragraph("• Vite")
doc.add_paragraph("• React Router DOM (v7)")
doc.add_paragraph("• QRCode.react")
doc.add_paragraph("• Vanilla CSS variables for custom theming")

# 5. Backend Technologies Used
add_heading(doc, '5. Backend Technologies Used', 1)
doc.add_paragraph("• Firebase Authentication: User session management.")
doc.add_paragraph("• Firebase Firestore: NoSQL document database (users, doctors, verificationKeys).")
doc.add_paragraph("• Firebase Storage: Secure bucket for uploaded medical documents.")

# 6. Authentication & Routing System
add_heading(doc, '6. Authentication & Routing System', 1)
doc.add_paragraph("The application protects routes via a `PrivateRoute` component in `App.jsx` which checks the user's role hydrated by `AuthContext`:")

add_code_block(doc, """// src/App.jsx
function PrivateRoute({ children, allowedRoles }) {
  const { currentUser, userData, loading } = useAuth();
  
  if (loading) return <div className="loading-page">Verifying session...</div>;
  if (!currentUser) return <Navigate to="/login" />;
  
  if (allowedRoles && userData) {
    if (!allowedRoles.includes(userData.role)) {
      // Redirect based on role
      return userData.role === 'doctor' 
        ? <Navigate to="/doctor-dashboard" /> 
        : <Navigate to="/dashboard" />;
    }
  }
  return children;
}""")

# 7. Database & Firebase Config
add_heading(doc, '7. Database & Firebase Config', 1)
doc.add_paragraph("The Firebase config utilizes environment variables for security, preventing API keys from leaking into the codebase:")

add_code_block(doc, """// src/firebase.js
import { initializeApp } from 'firebase/app';
import { getAuth } from 'firebase/auth';
import { getFirestore } from 'firebase/firestore';
import { getStorage } from 'firebase/storage';

const firebaseConfig = {
  apiKey: import.meta.env.VITE_FIREBASE_API_KEY,
  authDomain: import.meta.env.VITE_FIREBASE_AUTH_DOMAIN,
  projectId: import.meta.env.VITE_FIREBASE_PROJECT_ID,
  storageBucket: import.meta.env.VITE_FIREBASE_STORAGE_BUCKET,
  messagingSenderId: import.meta.env.VITE_FIREBASE_MESSAGING_SENDER_ID,
  appId: import.meta.env.VITE_FIREBASE_APP_ID
};

const app = initializeApp(firebaseConfig);
export const auth = getAuth(app);
export const db = getFirestore(app);
export const storage = getStorage(app);
export default app;""")

# 8. Doctor Verification Flow
add_heading(doc, '8. Doctor Verification Flow', 1)
doc.add_paragraph("The Verification Hub allows a doctor to fetch a document from a patient using a temporary key (e.g., AV-XXXX-XXXX):")

add_code_block(doc, """// src/pages/DoctorDashboard.jsx
async function validateKey(keyToValidate) {
  const q = query(
    collection(db, 'verificationKeys'),
    where('key', '==', keyToValidate.trim().toUpperCase())
  );
  const snap = await getDocs(q);

  const keyDoc = snap.docs[0];
  const keyData = { id: keyDoc.id, ...keyDoc.data() };

  // Check expiry
  if (new Date(keyData.expiresAt) <= new Date()) {
    await setDoc(doc(db, 'verificationKeys', keyDoc.id), 
      { status: 'expired' }, { merge: true });
    setKeyError('This verification key has expired.');
    return;
  }

  // Load patient & specific document
  const patientSnap = await getDoc(doc(db, 'users', keyData.patientId));
  const pData = patientSnap.data();
  const docIdx = keyData.documentIndex;
  
  if (pData.documents && pData.documents[docIdx]) {
    setTargetDocument(pData.documents[docIdx]);
    setAccessGranted(true);
  }
}""")

# 9. Emergency QR System
add_heading(doc, '9. Emergency QR System', 1)
doc.add_paragraph("The Emergency feature utilizes QRCode.react to dynamically generate a health card linking to a public route containing read-only critical medical data.")

add_code_block(doc, """// src/pages/EmergencyQRPage.jsx
import { QRCodeSVG } from 'qrcode.react';

// Generates URL linking to public emergency profile
const qrUrl = currentUser ? `${window.location.origin}/emergency/${currentUser.uid}` : '';

// Renders SVG QR code 
<QRCodeSVG 
  value={qrUrl} 
  size={140} 
  level="H" 
  includeMargin={true} 
  fgColor="#0b1f5f" 
  bgColor="#FFFFFF" 
/>""")

# 10. Document Generation & Context
add_heading(doc, '10. Theme Hydration (Context)', 1)
doc.add_paragraph("ArogyaVault uses `AuthContext` to fetch user profiles and globally inject CSS themes:")

add_code_block(doc, """// src/contexts/AuthContext.jsx
useEffect(() => {
  const unsubscribe = onAuthStateChanged(auth, async (user) => {
    if (user) {
      let docSnap = await getDoc(doc(db, 'users', user.uid));
      if (docSnap.exists()) {
        setUserData(docSnap.data());
        document.documentElement.setAttribute('data-theme', 'patient');
      } else {
        docSnap = await getDoc(doc(db, 'doctors', user.uid));
        if (docSnap.exists()) {
          setUserData(docSnap.data());
          document.documentElement.setAttribute('data-theme', 'doctor');
        }
      }
    }
  });
  return unsubscribe;
}, []);""")

# End of document
doc.save('ArogyaVault_Technical_Documentation.docx')
